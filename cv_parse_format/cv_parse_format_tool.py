"""CV Parser & Formatter — Cornerstone Project Source
Parses .docx/PDF CVs into a structured profile via the Claude API,
lets the consultant review/edit, then generates a formatted CV
(.docx + PDF) from a Word template with adjustable formatting.
"""

import base64
import json
import os
import re
import sys
import threading
import traceback
from datetime import datetime
from typing import List, Optional

import customtkinter as ctk
from tkinter import filedialog, messagebox

import anthropic
from pydantic import BaseModel, Field

import docx as pydocx
from docx.shared import Pt, Cm, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import pdfplumber


# ── First-run dependency bootstrap ─────────────────────────────────────────────
_BOOTSTRAP_STATUS = {"done": True, "ok": True, "msg": ""}

def _find_soffice():
    """Return path to a working LibreOffice binary, or None.

    Verifies the binary actually runs (--version) so stale symlinks or
    partially-uninstalled installs are not reported as available.
    Called fresh every time — no caching — so the gear panel always
    reflects the real current state.
    """
    import shutil, subprocess
    if sys.platform == "darwin":
        candidates = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            os.path.expanduser("~/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        ]
    elif sys.platform == "win32":
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
    else:
        candidates = []
    which = shutil.which("soffice")
    if which:
        candidates.insert(0, which)
    for c in candidates:
        if not c or not os.path.exists(c):
            continue
        # On Windows, file existence is sufficient — running soffice.exe
        # spawns a CMD window even with capture_output and may hang.
        # Stale-symlink detection only matters on macOS/Linux (Homebrew).
        if sys.platform == "win32":
            return c
        try:
            r = subprocess.run([c, "--version"], capture_output=True, timeout=8)
            if r.returncode == 0:
                return c
        except Exception:
            pass
    return None


def _bootstrap_dependencies():
    """Check what PDF tools are available and log it. No installs, no subprocesses."""
    if _find_soffice():
        print("[bootstrap] LibreOffice found ✓", flush=True)
    else:
        print("[bootstrap] LibreOffice not found — will use Word (docx2pdf) or PyMuPDF fallback", flush=True)


# ── Resource path (PyInstaller compat) ─────────────────────────────────────────
def resource_path(relative_path):
    try:
        base = sys._MEIPASS
    except Exception:
        base = os.path.abspath(".")
    return os.path.join(base, relative_path)


def app_dir():
    """User-writable data folder.
    Frozen: ~/Library/Application Support/Cornerstone Tools (macOS) or
            %APPDATA%\\Cornerstone Tools (Windows) — never inside the .app bundle
            which macOS App Translocation makes read-only.
    Dev:    the cv_parse_format/ source folder next to this script.
    """
    if getattr(sys, "frozen", False):
        if sys.platform == "win32":
            base = os.environ.get("APPDATA", os.path.expanduser("~"))
        else:
            base = os.path.join(os.path.expanduser("~"), "Library", "Application Support")
        d = os.path.join(base, "Cornerstone Tools")
        os.makedirs(d, exist_ok=True)
        return d
    return os.path.dirname(os.path.abspath(__file__))


def _config_path():
    """cv_config.json is bundled into sys._MEIPASS; fall back to app_dir for dev."""
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        p = os.path.join(meipass, "cv_config.json")
        if os.path.exists(p):
            return p
    return os.path.join(app_dir(), "cv_config.json")


CONFIG_PATH = _config_path()
PROFILES_DIR = os.path.join(app_dir(), "profiles")
OUTPUT_DIR = os.path.join(app_dir(), "output")
TEMPLATES_DIR = os.path.join(app_dir(), "templates")
DESIGN_DIR = os.path.join(app_dir(), "design")
DEFAULT_TEMPLATE = "Cornerstone.docx"


def _seed_bundled_assets():
    """Copy templates/design assets from sys._MEIPASS into app_dir on first run."""
    import shutil
    meipass = getattr(sys, "_MEIPASS", None)
    if not meipass:
        return
    for src_sub, dst_dir in [
        (os.path.join(meipass, "cv_parse_format", "templates"), TEMPLATES_DIR),
        (os.path.join(meipass, "cv_parse_format", "design"),    DESIGN_DIR),
    ]:
        if not os.path.isdir(src_sub):
            continue
        os.makedirs(dst_dir, exist_ok=True)
        for fname in os.listdir(src_sub):
            dst = os.path.join(dst_dir, fname)
            if not os.path.exists(dst):
                shutil.copy2(os.path.join(src_sub, fname), dst)

_seed_bundled_assets()

DEFAULT_FORMATTING = {
    "margin_top_cm": 4.4,
    "margin_bottom_cm": 2.4,
    "margin_left_cm": 1.27,
    "margin_right_cm": 1.27,
    "font_family": "Roboto",
    "template": "Cornerstone.docx",
    "name_size_pt": 20,
    "heading_size_pt": 12,
    "body_size_pt": 10,
    "accent_color": "#2a2a2a",
    "photo_enabled": False,
    "photo_width_mm": 35,
    "logo_path": "",
    "logo_width_mm": 40,
}


def _user_config_path():
    """Writable config path — never inside the read-only bundle."""
    return os.path.join(app_dir(), "cv_config.json")

def save_config(cfg):
    path = _user_config_path()
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w") as f:
        json.dump(cfg, f, indent=2)

def load_config():
    """Load config: bundled read-only defaults merged with user-writable overrides."""
    cfg = {}
    # 1. Bundled config (API keys baked in at build time)
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH) as f:
                cfg = json.load(f)
        except Exception:
            cfg = {}
    # 2. User-writable overlay (saved settings overrides)
    user_path = _user_config_path()
    if os.path.exists(user_path) and user_path != CONFIG_PATH:
        try:
            with open(user_path) as f:
                cfg.update(json.load(f))
        except Exception:
            pass
    fmt = dict(DEFAULT_FORMATTING)
    fmt.update(cfg.get("formatting", {}))
    cfg["formatting"] = fmt
    cfg.setdefault("anthropic_api_key", "")
    return cfg


CONFIG = load_config()

# ── Sheets store singleton ──────────────────────────────────────────────────────
_sheets_store_instance = None
_sheets_store_checked  = False

# ── Theme — light, paper-like ───────────────────────────────────────────────────
GOLD    = "#b8965a"   # accent (slightly deepened for light backgrounds)
GOLD_HV = "#a98549"
BG      = "#edeae3"   # warm paper background (slightly deeper so cards pop)
CARD    = "#ffffff"   # raised surface — pure white bubbles
SURFACE = "#f4f2ec"   # inputs / wells
HAIR    = "#e1dccf"   # hairline border
WHITE   = "#2a2a2a"   # primary text (dark ink on paper)
MUTED   = "#8d8779"   # secondary text
GREEN   = "#2e8f4e"
BLUE    = "#3a6fc4"
ORANGE  = "#bd7a1a"
RED     = "#bf4040"

# normal arrow cursor everywhere — no pointing hand on buttons
for _cls in ("CTkButton", "CTkCheckBox", "CTkSwitch", "CTkOptionMenu",
             "CTkSlider", "CTkSegmentedButton", "CTkRadioButton"):
    try:
        getattr(ctk, _cls)._set_cursor = lambda self: None
    except AttributeError:
        pass


def bind_hover(container, target, on_bg, off_bg, group=None):
    """Flicker-free hover: highlights `target` while the pointer is anywhere
    inside `container`. Pass a shared `group` dict for sibling rows so fast
    mouse movement can never leave a stale highlight behind."""
    path = str(container)

    def clear():
        target.configure(fg_color=off_bg)
        if group is not None and group.get("current") is target:
            group["current"] = None

    def on(_e=None):
        if group is not None:
            prev = group.get("current")
            if prev is not None and prev is not target:
                try:
                    prev.configure(fg_color=off_bg)
                except Exception:
                    pass
            group["current"] = target
        target.configure(fg_color=on_bg)

    def off(_e=None):
        x, y = container.winfo_pointerxy()
        w = container.winfo_containing(x, y)
        if w is None or not str(w).startswith(path):
            clear()

    def walk(wdg):
        wdg.bind("<Enter>", on, add="+")
        wdg.bind("<Leave>", off, add="+")
        for ch in wdg.winfo_children():
            walk(ch)
    walk(container)


def centered_column(parent, max_width=870):
    """A column that stays centred and never grows past max_width; it follows
    the window down when shrunk."""
    outer = ctk.CTkFrame(parent, fg_color="transparent")
    outer.pack(fill="both", expand=True)
    inner = ctk.CTkFrame(outer, fg_color="transparent", width=max_width)
    inner.place(relx=0.5, y=0, anchor="n", relheight=1.0)
    inner.pack_propagate(False)

    def resize(_e=None):
        inner.configure(width=max(min(outer.winfo_width() - 16, max_width), 320))
    outer.bind("<Configure>", resize)
    return inner


FONT      = ("Arial", 13)
FONT_SM   = ("Arial", 11)
FONT_BOLD = ("Arial", 13, "bold")
FONT_H1   = ("Arial", 20, "bold")


# ── Profile schema ─────────────────────────────────────────────────────────────
class WorkEntry(BaseModel):
    company: str = ""
    title: str = ""
    start: str = Field("", description="Start date as written on the CV, e.g. 'Mar 2021'")
    end: str = Field("", description="End date, or 'Present'")
    location: str = ""
    responsibilities: List[str] = []


class EducationEntry(BaseModel):
    institution: str = ""
    qualification: str = ""
    year: str = ""


class CandidateProfile(BaseModel):
    name: str = ""
    job_title: str = Field("", description="Current or most recent job title")
    email: str = ""
    phone: str = ""
    location: str = ""
    county: str = Field("", description="County or region (e.g. 'West Yorkshire', 'Greater Manchester')")
    linkedin: str = Field("", description="LinkedIn profile URL if present")
    availability: str = Field("", description="Notice period / availability if stated, e.g. '1 Month or less'")
    current_salary: str = Field("", description="Current salary if stated")
    desired_salary: str = Field("", description="Desired salary if stated, e.g. '£85,000 - £90,000'")
    right_to_work: str = Field("", description="Right to work / visa status if stated, e.g. 'Citizen'")
    motivations: str = Field("", description="Motivations for moving roles, if stated")
    summary: str = Field("", description="Professional summary / personal profile paragraph")
    skills: List[str] = []
    certifications: List[str] = Field([], description="Certifications, tickets, licences (e.g. CSCS, SMSTS, IPAF)")
    languages: List[str] = []
    achievements: List[str] = Field([], description="Notable achievements, awards, recognitions")
    interests: str = Field("", description="Hobbies / interests paragraph if present")
    work_history: List[WorkEntry] = []
    education: List[EducationEntry] = []


# ── Text extraction ────────────────────────────────────────────────────────────
def extract_docx_text(path):
    """Extract all text from a .docx regardless of layout complexity.

    Strategy: use python-docx XML directly to collect every w:t node's text.
    We walk the full XML tree so tables, text boxes, headers, footers, and
    shapes are all included. Paragraphs are separated by newlines; table cells
    by tabs so columns stay readable.
    """
    from docx.oxml.ns import qn
    doc = pydocx.Document(path)

    T   = qn("w:t")
    P   = qn("w:p")
    TR  = qn("w:tr")   # table row
    TC  = qn("w:tc")   # table cell

    parts = []

    def walk(node):
        """Recursively collect text, respecting paragraph/row boundaries."""
        for child in node:
            tag = child.tag
            if tag == TR:
                # Collect each cell's text, tab-separated, then newline
                row_parts = []
                for tc in child.iter(TC):
                    cell_text = " ".join(
                        "".join(t.text or "" for t in p.iter(T)).strip()
                        for p in tc.iter(P)
                        if "".join(t.text or "" for t in p.iter(T)).strip()
                    )
                    if cell_text:
                        row_parts.append(cell_text)
                if row_parts:
                    parts.append("\t".join(row_parts))
            elif tag == P:
                text = "".join(t.text or "" for t in child.iter(T)).strip()
                if text:
                    parts.append(text)
            else:
                walk(child)

    # Main document body
    walk(doc.element.body)

    # Also pull headers and footers (often contain name/contact details)
    for section in doc.sections:
        for hdr_ftr in [section.header, section.footer,
                        section.even_page_header, section.even_page_footer,
                        section.first_page_header, section.first_page_footer]:
            try:
                if hdr_ftr and hdr_ftr._element is not None:
                    walk(hdr_ftr._element)
            except Exception:
                pass

    text = "\n".join(parts)
    print(f"[extract_docx] extracted {len(text)} chars from {os.path.basename(path)}", flush=True)
    return text


def extract_doc_text(path):
    """Legacy .doc — convert to plain text. Uses textutil (built into macOS);
    on Windows asks the user to re-save as .docx."""
    import subprocess, tempfile
    if sys.platform != "darwin":
        raise RuntimeError("Legacy .doc files aren't supported on this machine — open it in Word and save as .docx first.")
    out = os.path.join(tempfile.mkdtemp(), "cv.txt")
    r = subprocess.run(["textutil", "-convert", "txt", "-output", out, path],
                       capture_output=True, text=True)
    if r.returncode != 0 or not os.path.exists(out):
        raise RuntimeError(f"Could not read this .doc file: {r.stderr.strip() or 'unknown error'}")
    with open(out, encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_pdf_text(path):
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t)
    return "\n".join(parts)


# ── Claude parsing ─────────────────────────────────────────────────────────────
PARSE_SYSTEM = (
    "You are a CV parsing assistant for a UK construction recruitment agency. "
    "Extract the candidate's details from the CV exactly as written — do not invent "
    "or embellish anything. Keep dates in the format they appear. If a field is not "
    "present in the CV, leave it empty. Summaries should be taken from the CV's own "
    "profile/summary section if present; otherwise write a brief factual 2-3 sentence "
    "summary of the candidate's experience. "
    "IMPORTANT: every employment/role mentioned anywhere in the CV must appear in "
    "work_history (employer, title, dates, location, responsibilities/duties as bullets) "
    "— even if the CV formats them unusually or lists them under 'Projects'/'Experience'. "
    "Project lists belong as responsibilities under the relevant job (or the most recent "
    "job if unattributed), NOT in achievements. All qualifications/degrees go in education. "
    "Use achievements only for awards and recognitions."
)


def find_claude_cli():
    import shutil
    found = shutil.which("claude")
    if found:
        return found
    for p in ("/opt/homebrew/bin/claude", "/usr/local/bin/claude",
              os.path.expanduser("~/.local/bin/claude"),
              os.path.expanduser("~/.claude/local/claude")):
        if os.path.exists(p):
            return p
    return None


def extract_profile_json(out):
    """Scan model output for a JSON object that validates as a CandidateProfile
    (the model may narrate before/after the JSON)."""
    print(f"[extract_profile_json] output preview: {out[:300]!r}", flush=True)
    decoder = json.JSONDecoder()
    profile = None
    i = out.find("{")
    while i != -1:
        try:
            obj, consumed = decoder.raw_decode(out[i:])
            if isinstance(obj, dict) and "name" in obj:
                try:
                    profile = CandidateProfile.model_validate(obj)
                    print(f"[extract_profile_json] validated: jobs={len(profile.work_history)} edu={len(profile.education)}", flush=True)
                except Exception as ve:
                    print(f"[extract_profile_json] validate failed: {ve}", flush=True)
            i = out.find("{", i + max(consumed, 1))
        except json.JSONDecodeError:
            i = out.find("{", i + 1)
    if profile is None:
        raise RuntimeError(f"Model returned no valid profile JSON:\n{out[:500]}")
    return profile


def parse_cv_via_claude_code(path):
    """Parse using the locally installed Claude Code CLI (covered by the
    Claude subscription — no API billing). Claude Code reads the file itself,
    so PDFs (including scanned) and .docx both work."""
    import subprocess
    cli = find_claude_cli()
    if not cli:
        raise RuntimeError(
            "No Anthropic API key set and Claude Code is not installed.\n"
            "Either add an API key in Settings, or install Claude Code "
            "(brew install claude-code) and sign in with your Claude subscription.")

    schema = json.dumps(CandidateProfile.model_json_schema())
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        text = extract_docx_text(path)
        if not text.strip():
            raise RuntimeError("Could not extract any text from this document.")
        task = f"Extract the candidate profile from this CV. Include EVERY job in work_history:\n\n{text}"
    elif ext == ".pdf":
        task = (f'Read the CV file at "{path}" and extract the candidate profile. '
                f'Include EVERY job in work_history — do not skip or merge any roles.')
    elif ext == ".doc":
        text = extract_doc_text(path)
        if not text.strip():
            raise RuntimeError("Could not extract any text from this document.")
        task = f"Extract the candidate profile from this CV. Include EVERY job in work_history:\n\n{text}"
    else:
        raise RuntimeError(f"Unsupported file type: {ext}")
    prompt = (
        f"{PARSE_SYSTEM}\n\n{task}\n\n"
        f"Respond with ONLY a single JSON object (no markdown fences, no commentary) "
        f"that validates against this JSON schema:\n{schema}"
    )
    try:
        result = subprocess.run(
            [cli, "-p", prompt, "--output-format", "text", "--allowedTools", "Read"],
            capture_output=True, text=True, timeout=300,
            cwd=os.path.dirname(path) or app_dir(),
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("Claude Code took too long to respond (5 min timeout).")
    out = (result.stdout or "").strip()
    if result.returncode != 0 or "Not logged in" in out or "/login" in out:
        raise RuntimeError(
            "Claude Code isn't signed in yet. Open Terminal, run `claude`, "
            "then type /login and sign in with your Claude account. "
            f"\n\nDetails: {out or result.stderr}".strip())
    return extract_profile_json(out)


def parse_cv(path):
    """Parse a CV file into a CandidateProfile. Uses the Anthropic API if a key
    is configured; otherwise falls back to the local Claude Code CLI (uses the
    Claude subscription, no separate billing). PDFs are sent natively (handles
    scanned documents); .docx is sent as extracted text."""
    # Reload config fresh each call so any settings changes take effect
    cfg = load_config()
    api_key = cfg.get("anthropic_api_key", "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    print(f"[parse_cv] config_path={CONFIG_PATH} api_key={'SET' if api_key else 'MISSING'}", flush=True)
    if not api_key:
        print("[parse_cv] No API key — falling back to Claude Code CLI", flush=True)
        return parse_cv_via_claude_code(path)
    client = anthropic.Anthropic(api_key=api_key)

    USER_INSTRUCTION = (
        "Extract the complete candidate profile from this CV. "
        "You MUST include EVERY job listed in the CV inside work_history — "
        "do not skip or merge any role, no matter how many there are. "
        "Respond with ONLY a single JSON object (no markdown fences, no commentary) "
        "matching this schema:\n"
    )

    ext = os.path.splitext(path)[1].lower()
    schema = json.dumps(CandidateProfile.model_json_schema(), indent=None)
    instruction_block = USER_INSTRUCTION + schema

    def _doc_as_pdf_b64(file_path):
        """Convert DOCX/DOC → PDF → base64 for the Anthropic API via LibreOffice."""
        import tempfile, subprocess
        soffice = _find_soffice()
        if not soffice:
            raise RuntimeError(
                "LibreOffice is not installed.\n\n"
                "Please install it using the ⚙ Tools button in the sidebar, "
                "then try again."
            )
        with tempfile.TemporaryDirectory() as tmp:
            kw = {"capture_output": True, "text": True, "timeout": 120}
            if sys.platform == "win32":
                si = subprocess.STARTUPINFO()
                si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                si.wShowWindow = subprocess.SW_HIDE
                kw["startupinfo"] = si
                kw["creationflags"] = subprocess.CREATE_NO_WINDOW
            r = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, file_path],
                **kw)
            lo_pdf = os.path.join(
                tmp, os.path.splitext(os.path.basename(file_path))[0] + ".pdf")
            if r.returncode == 0 and os.path.exists(lo_pdf):
                with open(lo_pdf, "rb") as f:
                    data = f.read()
                print(f"[parse_cv] LibreOffice DOCX→PDF OK ({len(data)//1024}KB)", flush=True)
                return base64.standard_b64encode(data).decode()
            raise RuntimeError(
                f"LibreOffice failed to prepare the CV for parsing.\n\n"
                f"Details: {r.stderr[:300] or 'no output'}"
            )

    def _make_doc_content(pdf_b64):
        return [
            {"type": "document",
             "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64}},
            {"type": "text", "text": instruction_block},
        ]

    if ext == ".pdf":
        with open(path, "rb") as f:
            pdf_b64 = base64.standard_b64encode(f.read()).decode()
        content = _make_doc_content(pdf_b64)
        print(f"[parse_cv] sending PDF directly ({len(pdf_b64)//1024}KB)", flush=True)

    elif ext in (".docx", ".doc"):
        # Convert to PDF via PyMuPDF so the API sees the whole rendered document —
        # tables, text boxes, columns, all formatting — not just extracted text.
        try:
            pdf_b64 = _doc_as_pdf_b64(path)
            content = _make_doc_content(pdf_b64)
            print(f"[parse_cv] converted {ext} → PDF ({len(pdf_b64)//1024}KB)", flush=True)
        except Exception as conv_err:
            # Fallback: extract text (better than nothing)
            print(f"[parse_cv] PDF conversion failed ({conv_err}), falling back to text extraction", flush=True)
            text = extract_docx_text(path) if ext == ".docx" else extract_doc_text(path)
            if not text.strip():
                raise RuntimeError("Could not extract any text from this document.")
            content = f"{instruction_block}\n\n---\n{text}"

    else:
        raise RuntimeError(f"Unsupported file type: {ext}")

    model = cfg.get("parse_model", "claude-sonnet-4-6")
    print(f"[parse_cv] calling Anthropic API model={model} ext={ext}", flush=True)
    try:
        response = client.messages.create(
            model=model,
            max_tokens=16000,
            system=PARSE_SYSTEM,
            messages=[{"role": "user", "content": content}],
        )
    except Exception as api_err:
        print(f"[parse_cv] API error: {api_err}", flush=True)
        raise
    out = "".join(b.text for b in response.content if b.type == "text")
    print(f"[parse_cv] response length={len(out)} stop_reason={response.stop_reason}", flush=True)
    return extract_profile_json(out)


# ── Default template generation ────────────────────────────────────────────────
BODY_GREY = RGBColor(0x56, 0x56, 0x54)   # body text grey from the sample CVs
DARK      = RGBColor(0x2A, 0x2A, 0x2A)   # headings near-black from the sample CVs


def list_templates():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    return sorted(f for f in os.listdir(TEMPLATES_DIR)
                  if f.endswith(".docx") and not f.startswith("~"))


def template_path(name=None):
    return os.path.join(TEMPLATES_DIR, name or CONFIG["formatting"].get("template", DEFAULT_TEMPLATE))


def ensure_template():
    """Create the default house template (templates/Cornerstone.docx) if missing.
    The template is a normal Word file with {{ placeholders }} — open it in Word
    to change the design, or drop alternative .docx templates into templates/
    and pick them in the app. Header/footer artwork comes from design/header.png
    and design/footer.png. Delete the template file to regenerate."""
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    tpl_file = os.path.join(TEMPLATES_DIR, DEFAULT_TEMPLATE)
    if os.path.exists(tpl_file):
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

    doc = pydocx.Document()
    usable_cm = 21.0 - 1.27 * 2  # A4 width minus side margins

    # Page geometry — generous top/bottom margins clear the artwork bands
    for section in doc.sections:
        section.top_margin = Cm(4.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.3)

    # House artwork in header & footer (repeats on every page)
    hdr_img = os.path.join(DESIGN_DIR, "header.png")
    ftr_img = os.path.join(DESIGN_DIR, "footer.png")
    if os.path.exists(hdr_img):
        hp = doc.sections[0].header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.add_run().add_picture(hdr_img, width=Cm(usable_cm))
    if os.path.exists(ftr_img):
        fp = doc.sections[0].footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run().add_picture(ftr_img, width=Cm(usable_cm))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Roboto"
    normal.font.size = Pt(10)
    normal.font.color.rgb = BODY_GREY
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Roboto"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = DARK
    title.paragraph_format.space_after = Pt(14)
    # subtle hairline rule under the job title (replaces Word's heavy default)
    from docx.oxml.ns import qn
    pPr = title.element.get_or_add_pPr()
    for bdr in pPr.findall(qn("w:pBdr")):
        pPr.remove(bdr)
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {})
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "C8C2B6")
    pbdr.append(bottom)
    pPr.append(pbdr)

    h1 = styles["Heading 1"]
    h1.font.name = "Roboto"
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = DARK
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.name = "Roboto"
    h2.font.size = Pt(10)
    h2.font.bold = True
    h2.font.color.rgb = DARK
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.tab_stops.add_tab_stop(Cm(usable_cm), WD_TAB_ALIGNMENT.RIGHT)

    bullet = styles["List Bullet"]
    bullet.font.name = "Roboto"
    bullet.font.size = Pt(10)
    bullet.font.color.rgb = BODY_GREY
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.left_indent = Cm(0.55)
    bullet.paragraph_format.first_line_indent = Cm(-0.3)

    def fact_row(label, var):
        doc.add_paragraph(f"{{%p if {var} %}}")
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(4.4), WD_TAB_ALIGNMENT.LEFT)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10.8)
        r.font.color.rgb = DARK
        p.add_run(f"\t{{{{ {var} }}}}")
        doc.add_paragraph("{%p endif %}")

    # Optional extra logo image (Settings → logo path); the main branding
    # lives in the header artwork above
    doc.add_paragraph("{%p if logo %}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{ logo }}")
    doc.add_paragraph("{%p endif %}")

    # Anonymised header: job title only
    doc.add_paragraph("{{ job_title }}", style="Title")

    fact_row("Availability", "availability")
    fact_row("Desired Salary", "desired_salary")
    fact_row("Location", "location")

    doc.add_paragraph("{%p if summary %}")
    p = doc.add_paragraph("{{ summary }}")
    p.paragraph_format.space_before = Pt(12)
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if work_history %}")
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("{%p for job in work_history %}")
    doc.add_paragraph(
        "{{ job.company }}{% if job.location %}, {{ job.location }}{% endif %}"
        "{% if job.start %}\t{{ job.start }} - {{ job.end }}{% endif %}",
        style="Heading 2")
    doc.add_paragraph("{%p if job.title %}")
    doc.add_paragraph("{{ job.title }}").runs[0].italic = True
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p for r in job.responsibilities %}")
    doc.add_paragraph("{{ r }}", style="List Bullet")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if education %}")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("{%p for e in education %}")
    doc.add_paragraph("{{ e.institution }}{% if e.year %}\t{{ e.year }}{% endif %}", style="Heading 2")
    doc.add_paragraph("{{ e.qualification }}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if skills %}")
    doc.add_heading("Skills", level=1)
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "{%tr for pair in skill_pairs %}"
    tbl.cell(1, 0).text = "{{ pair[0] }}"
    tbl.cell(1, 1).text = "{{ pair[1] }}"
    tbl.cell(2, 0).text = "{%tr endfor %}"
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if certifications %}")
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("{%p for c in certifications %}")
    doc.add_paragraph("{{ c }}", style="List Bullet")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if achievements %}")
    doc.add_heading("Achievements", level=1)
    doc.add_paragraph("{%p for a in achievements %}")
    doc.add_paragraph("{{ a }}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if interests %}")
    doc.add_heading("Interests", level=1)
    doc.add_paragraph("{{ interests }}")
    doc.add_paragraph("{%p endif %}")

    doc.save(tpl_file)


# ── Formatted CV generation ────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def fix_theme_fonts(docx_path):
    """Strip theme-font attributes from styles so the template's own explicit
    fonts render true (theme attrs override them and cause Carlito/Calibri
    substitution in LibreOffice/Word). No design overrides — the .docx
    template fully controls margins, fonts, sizes and colours."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    doc = pydocx.Document(docx_path)
    for st in doc.styles:
        try:
            rPr = st.element.get_or_add_rPr()
        except (AttributeError, ValueError):
            continue
        rFonts = rPr.find(W + "rFonts")
        if rFonts is None:
            continue
        explicit = rFonts.get(W + "ascii")
        if explicit:
            for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
                rFonts.attrib.pop(W + attr, None)
            rFonts.set(W + "hAnsi", explicit)
    doc.save(docx_path)


def convert_to_pdf(docx_path, pdf_path):
    """Convert DOCX → PDF using LibreOffice headless.

    LibreOffice is required — no fallbacks. If it is not installed the user
    is shown a clear message via the warning banner and the ⚙ Tools button.
    """
    import subprocess

    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice is not installed.\n\n"
            "Please install it using the ⚙ Tools button in the sidebar, "
            "then try again."
        )

    out_dir = os.path.dirname(pdf_path) or "."
    kw = {"capture_output": True, "text": True, "timeout": 120}
    if sys.platform == "win32":
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        si.wShowWindow = subprocess.SW_HIDE
        kw["startupinfo"] = si
        kw["creationflags"] = subprocess.CREATE_NO_WINDOW
    r = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        **kw)
    lo_pdf = os.path.join(out_dir,
                          os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if r.returncode == 0 and os.path.exists(lo_pdf):
        if lo_pdf != pdf_path:
            os.replace(lo_pdf, pdf_path)
        print(f"[convert_to_pdf] LibreOffice OK: {os.path.basename(pdf_path)}", flush=True)
        return
    raise RuntimeError(
        f"LibreOffice failed to convert the CV to PDF.\n\n"
        f"Details: {r.stderr[:300] or 'no output'}"
    )


def generate_cv(profile: CandidateProfile, fmt, photo_path=None, keep_docx=False, template_name=None, out_dir=None):
    ensure_template()
    out_dir = out_dir or OUTPUT_DIR
    os.makedirs(out_dir, exist_ok=True)

    tpl_path = template_path(template_name)
    if not os.path.exists(tpl_path):
        raise RuntimeError(f"Template not found: {tpl_path}")
    tpl = DocxTemplate(tpl_path)
    ctx = profile.model_dump()
    ctx["facts"] = [{"label": lbl, "value": val} for lbl, val in (
        ("Availability", profile.availability),
        ("Desired Salary", profile.desired_salary),
        ("Location", profile.location),
    ) if val.strip()]
    skills = list(profile.skills)
    ctx["skill_pairs"] = [(skills[i], skills[i + 1] if i + 1 < len(skills) else "")
                          for i in range(0, len(skills), 2)]
    logo = fmt.get("logo_path", "").strip()
    if logo and os.path.exists(logo):
        ctx["logo"] = InlineImage(tpl, logo, width=Mm(fmt.get("logo_width_mm", 40)))
    else:
        ctx["logo"] = ""
    if photo_path and fmt.get("photo_enabled"):
        ctx["photo"] = InlineImage(tpl, photo_path, width=Mm(fmt["photo_width_mm"]))
    else:
        ctx["photo"] = ""
    tpl.render(ctx, autoescape=True)

    safe_name = re.sub(r"[^\w\- ]", "", profile.name or "Candidate").strip().replace(" ", "_")
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    docx_out = os.path.join(out_dir, f"{safe_name}_CV_{stamp}.docx")
    tpl.save(docx_out)
    fix_theme_fonts(docx_out)

    pdf_out = docx_out.replace(".docx", ".pdf")
    convert_to_pdf(docx_out, pdf_out)
    if not keep_docx:
        try:
            os.remove(docx_out)
        except OSError:
            pass
        docx_out = None
    return docx_out, pdf_out


# ── Profile persistence ────────────────────────────────────────────────────────
# v2 format: {"schema": 2, "profile": {...}, "cv": {...}}
# profile = parsed/corrected candidate record (the backup / source of truth)
# cv      = independent copy used for the formatted CV only

# ── Profile store (Google Sheets or local JSON fallback) ───────────────────────

def _get_sheets_store():
    """
    Return a cached SheetsStore if configured and reachable, else None.
    Connection is attempted once; subsequent calls reuse the same instance.
    """
    global _sheets_store_instance, _sheets_store_checked
    if _sheets_store_checked:
        return _sheets_store_instance
    _sheets_store_checked = True
    try:
        # Re-read config from disk so the google_sheets key is always fresh
        cfg = load_config()
        from google_sheets_store import from_config
        store = from_config(cfg)
        if store and store.is_available():
            _sheets_store_instance = store
            print("[sheets] Connected to Google Sheets ✓")
        else:
            print("[sheets] Not configured — using local JSON fallback")
    except Exception as e:
        print(f"[sheets] Connection failed: {e}")
    return _sheets_store_instance


def save_profile(profile: CandidateProfile, cv: Optional[CandidateProfile] = None,
                 existing_path=None, raw_cv_link: str = ""):
    """
    Save to Google Sheets when available, otherwise fall back to local JSON.
    Returns an opaque 'path' string (numeric profile_id for Sheets, file path for local).
    cv defaults to a copy of profile when the CV form hasn't been populated yet.
    """
    profile_dict = profile.model_dump()
    # If CV tab hasn't been filled yet, seed cv from profile so cv_json is never blank
    cv_dict = cv.model_dump() if cv else profile_dict.copy()

    store = _get_sheets_store()
    if store:
        # Only pass a numeric existing ID; let Sheets generate a new one otherwise
        existing_id = None
        if existing_path:
            try:
                int(str(existing_path))   # valid numeric ID?
                existing_id = str(existing_path)
            except (ValueError, TypeError):
                pass
        return store.save_profile(existing_id, profile_dict, cv_dict, raw_cv_link)

    # ── local JSON fallback ──
    os.makedirs(PROFILES_DIR, exist_ok=True)
    if existing_path and os.path.exists(str(existing_path)):
        path = existing_path
    else:
        safe_name = re.sub(r"[^\w\- ]", "", profile.name or "Candidate").strip().replace(" ", "_")
        path = os.path.join(PROFILES_DIR, f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
    data = {"schema": 2, "profile": profile_dict, "cv": cv_dict}
    with open(path, "w") as f:
        json.dump(data, f, indent=2)
    return path


def load_profile(profile_ref):
    """
    Returns (profile, cv_or_None).
    profile_ref is either a profile_id (Sheets) or a filename (local).
    """
    store = _get_sheets_store()
    if store and (not os.sep in str(profile_ref)):
        data = store.load_profile(profile_ref)
        if data:
            profile = CandidateProfile.model_validate(data["profile"])
            cv      = CandidateProfile.model_validate(data["cv"]) if data.get("cv") else None
            return profile, cv

    # ── local JSON fallback ──
    path = profile_ref if os.path.isabs(str(profile_ref)) \
           else os.path.join(PROFILES_DIR, profile_ref)
    with open(path) as f:
        data = json.load(f)
    if isinstance(data, dict) and data.get("schema") == 2:
        profile = CandidateProfile.model_validate(data["profile"])
        cv      = CandidateProfile.model_validate(data["cv"]) if data.get("cv") else None
        return profile, cv
    return CandidateProfile.model_validate(data), None


def list_profiles():
    """
    Return a list of profile summaries.
    Each item is a dict with at minimum: profile_id/filename, name, job_title,
    parsed_date, work_count, edu_count.
    """
    store = _get_sheets_store()
    if store:
        return store.list_profiles()   # already sorted newest-first

    # ── local JSON fallback ──
    if not os.path.isdir(PROFILES_DIR):
        return []
    files = sorted(
        (f for f in os.listdir(PROFILES_DIR) if f.endswith(".json")),
        key=lambda f: os.path.getmtime(os.path.join(PROFILES_DIR, f)),
        reverse=True,
    )
    summaries = []
    for fn in files:
        fp = os.path.join(PROFILES_DIR, fn)
        try:
            with open(fp) as fh:
                d = json.load(fh)
            pr = d.get("profile", d)
            summaries.append({
                "profile_id":  fn,
                "name":        pr.get("name", ""),
                "job_title":   pr.get("job_title", ""),
                "email":       pr.get("email", ""),
                "parsed_date": datetime.fromtimestamp(
                                   os.path.getmtime(fp)).strftime("%Y-%m-%d %H:%M"),
                "work_count":  len(pr.get("work_history") or []),
                "edu_count":   len(pr.get("education") or []),
            })
        except Exception:
            continue
    return summaries


def _new_profile_id(profile: CandidateProfile) -> str:
    safe = re.sub(r"[^\w\-]", "", (profile.name or "candidate").replace(" ", "_"))
    return f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"


# ── UI ─────────────────────────────────────────────────────────────────────────


class ProfileForm:
    """Reusable candidate form: bubbles + collapsible, reorderable cards.
    Used twice — Profile tab (source of truth) and CV tab (formatted-CV copy)."""

    PERSONAL = ("name", "linkedin", "email", "phone", "current_salary",
                "right_to_work", "motivations")

    def __init__(self, app, parent, mode="profile", on_change=None):
        self.app = app
        self.mode = mode          # "profile" = full record; "cv" = CV fields only
        self.on_change = on_change
        self.scroll = ctk.CTkScrollableFrame(parent, fg_color=BG)
        self.scroll.pack(fill="both", expand=True)
        self.widgets = {}
        self.job_cards = []
        self.edu_cards = []
        self.populated = False
        self._base = CandidateProfile()

    # ── primitives ──
    def _changed(self, _e=None):
        if self.on_change:
            self.on_change()

    def _entry(self, parent, label, value, multiline=False):
        ctk.CTkLabel(parent, text=label, font=FONT_SM, text_color=MUTED
                     ).pack(anchor="w", padx=12, pady=(8, 1))
        if multiline:
            w = ctk.CTkTextbox(parent, fg_color="#fbfaf7", border_color=HAIR, border_width=1,
                               corner_radius=8, text_color=WHITE, height=88, wrap="word")
            w.insert("1.0", value)
        else:
            w = ctk.CTkEntry(parent, fg_color="#fbfaf7", border_color=HAIR, border_width=1,
                             corner_radius=8, text_color=WHITE, height=32)
            w.insert(0, value)
        w.pack(fill="x", padx=12)
        w.bind("<KeyRelease>", self._changed)
        return w

    def _bubble(self, parent):
        card = ctk.CTkFrame(parent, fg_color=CARD, corner_radius=14,
                            border_width=1, border_color=HAIR)
        card.pack(fill="x", padx=10, pady=8)
        return card

    @staticmethod
    def _get(w):
        if isinstance(w, ctk.CTkTextbox):
            return w.get("1.0", "end").strip()
        return w.get().strip()

    @staticmethod
    def _get_lines(w):
        return [l.strip() for l in w.get("1.0", "end").splitlines() if l.strip()]

    # ── collapsible cards ──
    def _make_card(self, container, cards_list, title_fn, build_body):
        card = ctk.CTkFrame(container, fg_color=CARD, corner_radius=14,
                            border_width=1, border_color=HAIR)
        card.pack(fill="x", padx=10, pady=6)
        rec = {"frame": card, "collapsed": True, "widgets": {}}

        header = ctk.CTkFrame(card, fg_color="transparent", corner_radius=10)
        header.pack(fill="x", padx=8, pady=6)
        handle = ctk.CTkLabel(header, text="⠿", font=FONT, text_color=MUTED, width=18)
        handle.pack(side="left", padx=(4, 4))

        # Create tbox but DON'T pack yet — right buttons must be packed first
        tbox = ctk.CTkFrame(header, fg_color="transparent")
        rec["title_lbl"] = ctk.CTkLabel(tbox, text="", font=FONT_BOLD, text_color=WHITE, anchor="w")
        rec["title_lbl"].pack(anchor="w")
        rec["sub_lbl"] = ctk.CTkLabel(tbox, text="", font=FONT_SM, text_color=MUTED, anchor="w")
        rec["sub_lbl"].pack(anchor="w")

        def btn(txt, cmd, color=MUTED):
            return ctk.CTkButton(header, text=txt, width=28, height=26, fg_color="transparent",
                                 hover_color=SURFACE, text_color=color, command=cmd)

        # Pack right buttons BEFORE tbox — otherwise expand=True leaves no room for them.
        # Reverse visual order: last packed = leftmost (closest to tbox).
        btn("🗑", lambda: self._card_delete(rec, cards_list)).pack(side="right", padx=1)
        btn("↓", lambda: self._card_move(rec, cards_list, 1)).pack(side="right", padx=1)
        btn("↑", lambda: self._card_move(rec, cards_list, -1)).pack(side="right", padx=1)
        rec["chevron"] = ctk.CTkButton(header, text="▾", width=28, height=26,
                                       fg_color="transparent", hover_color=SURFACE,
                                       text_color=GOLD, command=lambda: self._card_toggle(rec),
                                       font=ctk.CTkFont("Arial", 18))
        rec["chevron"].pack(side="right", padx=1)

        # Now pack tbox — fills all remaining left space
        tbox.pack(side="left", fill="x", expand=True)

        rec["body"] = ctk.CTkFrame(card, fg_color="transparent")
        build_body(rec)
        rec["title_fn"] = title_fn

        def refresh_header(_=None):
            t, sub = title_fn(rec["widgets"])
            rec["title_lbl"].configure(text=t or "Untitled")
            rec["sub_lbl"].configure(text=sub)
        rec["refresh_header"] = refresh_header
        refresh_header()
        for w in rec["widgets"].values():
            if isinstance(w, ctk.CTkEntry):
                w.bind("<KeyRelease>", refresh_header, add="+")
        # Bind the whole header row (including tbox area) to toggle expand/collapse
        for clickable in (header, tbox, rec["title_lbl"], rec["sub_lbl"], handle):
            clickable.bind("<Button-1>", lambda _e, r=rec: self._card_toggle(r))
        bind_hover(header, header, SURFACE, "transparent")
        cards_list.append(rec)
        return rec

    def _card_toggle(self, rec):
        rec["collapsed"] = not rec["collapsed"]
        if rec["collapsed"]:
            rec["body"].pack_forget()
            rec["chevron"].configure(text="▾")
            rec["refresh_header"]()
        else:
            rec["body"].pack(fill="x", padx=8, pady=(0, 10))
            rec["chevron"].configure(text="▴")

    def _card_move(self, rec, cards_list, delta):
        i = cards_list.index(rec)
        j = i + delta
        if not (0 <= j < len(cards_list)):
            return
        cards_list[i], cards_list[j] = cards_list[j], cards_list[i]
        for r in cards_list:
            r["frame"].pack_forget()
        for r in cards_list:
            r["frame"].pack(fill="x", padx=10, pady=6)
        self._changed()

    def _card_delete(self, rec, cards_list):
        if not messagebox.askyesno("Remove entry", "Remove this entry?"):
            return
        cards_list.remove(rec)
        rec["frame"].destroy()
        self._changed()

    def _job_card(self, job):
        def build_body(rec):
            b = rec["body"]
            rec["widgets"]["title"] = self._entry(b, "Job title", job.title)
            rec["widgets"]["company"] = self._entry(b, "Company", job.company)
            row = ctk.CTkFrame(b, fg_color="transparent")
            row.pack(fill="x")
            h1 = ctk.CTkFrame(row, fg_color="transparent"); h1.pack(side="left", fill="x", expand=True)
            h2 = ctk.CTkFrame(row, fg_color="transparent"); h2.pack(side="left", fill="x", expand=True)
            rec["widgets"]["start"] = self._entry(h1, "Start date", job.start)
            rec["widgets"]["end"] = self._entry(h2, "End date", job.end)
            rec["widgets"]["location"] = self._entry(b, "Location", job.location)
            rec["widgets"]["responsibilities"] = self._entry(
                b, "Description (one bullet per line)", "\n".join(job.responsibilities), multiline=True)

        def title_fn(w):
            t, c = self._get(w["title"]), self._get(w["company"])
            dates = " - ".join(x for x in (self._get(w["start"]), self._get(w["end"])) if x)
            return " - ".join(x for x in (t, c) if x), dates
        self._make_card(self.jobs_container, self.job_cards, title_fn, build_body)

    def _edu_card(self, e):
        def build_body(rec):
            b = rec["body"]
            rec["widgets"]["qualification"] = self._entry(b, "Qualification", e.qualification)
            rec["widgets"]["institution"] = self._entry(b, "Institution", e.institution)
            rec["widgets"]["year"] = self._entry(b, "Year", e.year)

        def title_fn(w):
            q, i = self._get(w["qualification"]), self._get(w["institution"])
            return " - ".join(x for x in (q, i) if x), self._get(w["year"])
        self._make_card(self.edu_container, self.edu_cards, title_fn, build_body)

    # ── populate / collect ──
    def populate(self, p: CandidateProfile):
        self._base = p
        for w in self.scroll.winfo_children():
            w.destroy()
        self.widgets = {}
        self.job_cards = []
        self.edu_cards = []
        rs = self.scroll
        w = self.widgets

        details = self._bubble(rs)
        ctk.CTkLabel(details, text="Candidate Details" if self.mode == "profile" else "CV Header",
                     font=FONT_BOLD, text_color=GOLD).pack(anchor="w", padx=12, pady=(12, 0))
        grid = ctk.CTkFrame(details, fg_color="transparent")
        grid.pack(fill="x", pady=(0, 4))
        colL = ctk.CTkFrame(grid, fg_color="transparent"); colL.pack(side="left", fill="x", expand=True)
        colR = ctk.CTkFrame(grid, fg_color="transparent"); colR.pack(side="left", fill="x", expand=True)
        if self.mode == "profile":
            w["name"] = self._entry(colL, "Name", p.name)
            w["linkedin"] = self._entry(colR, "LinkedIn", p.linkedin)
            w["job_title"] = self._entry(colL, "Job title", p.job_title)
            w["location"] = self._entry(colR, "Location", p.location)
            w["email"] = self._entry(colL, "Email", p.email)
            w["phone"] = self._entry(colR, "Phone", p.phone)
            w["current_salary"] = self._entry(colL, "Current salary", p.current_salary)
            w["desired_salary"] = self._entry(colR, "Desired salary", p.desired_salary)
            w["availability"] = self._entry(colL, "Availability", p.availability)
            w["right_to_work"] = self._entry(colR, "Right to work", p.right_to_work)
            w["motivations"] = self._entry(details, "Motivations for moving", p.motivations, multiline=True)
        else:
            # CV copy — anonymised: no personal/contact details here
            w["job_title"] = self._entry(colL, "Job title (CV header)", p.job_title)
            w["location"] = self._entry(colR, "Location", p.location)
            w["availability"] = self._entry(colL, "Availability", p.availability)
            w["desired_salary"] = self._entry(colR, "Desired salary", p.desired_salary)
        w["summary"] = self._entry(details, "Profile summary", p.summary, multiline=True)
        ctk.CTkFrame(details, fg_color="transparent", height=8).pack()

        ctk.CTkLabel(rs, text="Work Experience", font=FONT_BOLD, text_color=WHITE
                     ).pack(anchor="w", padx=12, pady=(14, 0))
        self.jobs_container = ctk.CTkFrame(rs, fg_color="transparent")
        self.jobs_container.pack(fill="x")
        for job in p.work_history:
            self._job_card(job)
        ctk.CTkButton(rs, text="+ Add Experience", fg_color="transparent", hover_color=SURFACE,
                      text_color=BLUE, anchor="w",
                      command=lambda: (self._job_card(WorkEntry()), self._changed())
                      ).pack(anchor="w", padx=12)

        ctk.CTkLabel(rs, text="Education", font=FONT_BOLD, text_color=WHITE
                     ).pack(anchor="w", padx=12, pady=(14, 0))
        self.edu_container = ctk.CTkFrame(rs, fg_color="transparent")
        self.edu_container.pack(fill="x")
        for e in p.education:
            self._edu_card(e)
        ctk.CTkButton(rs, text="+ Add Education", fg_color="transparent", hover_color=SURFACE,
                      text_color=BLUE, anchor="w",
                      command=lambda: (self._edu_card(EducationEntry()), self._changed())
                      ).pack(anchor="w", padx=12)

        extras = self._bubble(rs)
        ctk.CTkLabel(extras, text="Skills & Extras", font=FONT_BOLD,
                     text_color=GOLD).pack(anchor="w", padx=12, pady=(12, 0))
        w["skills"] = self._entry(extras, "Skills (one per line)", "\n".join(p.skills), multiline=True)
        w["certifications"] = self._entry(extras, "Certifications & Tickets (one per line)",
                                          "\n".join(p.certifications), multiline=True)
        w["languages"] = self._entry(extras, "Languages (one per line)", "\n".join(p.languages), multiline=True)
        w["achievements"] = self._entry(extras, "Achievements (one per line)",
                                        "\n".join(p.achievements), multiline=True)
        w["interests"] = self._entry(extras, "Interests", p.interests, multiline=True)
        ctk.CTkFrame(extras, fg_color="transparent", height=8).pack()
        ctk.CTkFrame(rs, fg_color="transparent", height=10).pack()
        self.populated = True

    def collect(self) -> Optional[CandidateProfile]:
        if not self.populated:
            return None
        w = self.widgets
        b = self._base

        def val(key, fallback):
            return self._get(w[key]) if key in w else fallback
        return CandidateProfile(
            name=val("name", b.name), job_title=val("job_title", b.job_title),
            email=val("email", b.email), phone=val("phone", b.phone),
            location=val("location", b.location), linkedin=val("linkedin", b.linkedin),
            availability=val("availability", b.availability),
            current_salary=val("current_salary", b.current_salary),
            desired_salary=val("desired_salary", b.desired_salary),
            right_to_work=val("right_to_work", b.right_to_work),
            motivations=val("motivations", b.motivations),
            summary=val("summary", b.summary),
            skills=self._get_lines(w["skills"]),
            certifications=self._get_lines(w["certifications"]),
            languages=self._get_lines(w["languages"]),
            achievements=self._get_lines(w["achievements"]),
            interests=val("interests", b.interests),
            work_history=[WorkEntry(
                company=self._get(j["widgets"]["company"]), title=self._get(j["widgets"]["title"]),
                start=self._get(j["widgets"]["start"]), end=self._get(j["widgets"]["end"]),
                location=self._get(j["widgets"]["location"]),
                responsibilities=self._get_lines(j["widgets"]["responsibilities"]),
            ) for j in self.job_cards],
            education=[EducationEntry(
                institution=self._get(e["widgets"]["institution"]),
                qualification=self._get(e["widgets"]["qualification"]),
                year=self._get(e["widgets"]["year"]),
            ) for e in self.edu_cards],
        )


CHIP_COLORS = ["#e8f3ec", "#fdeef0", "#eef1fb", "#fbf3e4", "#f1eafa", "#e9f5f8"]
CHIP_TEXT   = ["#2e7d4f", "#b04a5a", "#3a5bb0", "#a06b1a", "#6b3fa0", "#1a7d96"]


def time_ago(ts):
    secs = (datetime.now() - datetime.fromtimestamp(ts)).total_seconds()
    for limit, div, unit in ((120, 60, "a minute ago"), (7200, 3600, "an hour ago"),
                             (172800, 86400, "a day ago")):
        if secs < limit:
            return unit
    if secs < 3600:
        return f"{int(secs // 60)} minutes ago"
    if secs < 86400:
        return f"{int(secs // 3600)} hours ago"
    return f"{int(secs // 86400)} days ago"


class CVParseFormatTool(ctk.CTkFrame):
    """Embeddable frame — hosted inside Cornerstone Tools, or standalone via __main__."""

    def __init__(self, master=None):
        super().__init__(master, fg_color=BG)
        self.profile_path = None           # currently loaded/saved profile ref
        self._pending_raw_cv_link = ""     # Drive link for the just-parsed CV file
        self.photo_path = None
        self.preview_pil = []          # original PIL pages
        self.preview_pages = []        # CTkImages scaled to pane
        self.preview_idx = 0
        self._preview_after = None
        self._preview_busy = False
        self._preview_dirty = False
        self._build()
        _bootstrap_dependencies()

    def _build(self):
        main = self
        self._main = main

        # ── bootstrap-style navbar: full-width bar, brand left, inline links ──
        navbar = ctk.CTkFrame(main, fg_color=CARD, corner_radius=0, height=58)
        navbar.pack(fill="x")
        navbar.pack_propagate(False)
        ctk.CTkFrame(main, fg_color=HAIR, height=1).pack(fill="x")

        ctk.CTkLabel(navbar, text="CV Parse & Format Tool", font=("Arial", 17, "bold"),
                     text_color=WHITE).pack(side="left", padx=(24, 4))
        ctk.CTkLabel(navbar, text="V0.1", font=FONT_SM,
                     text_color=MUTED).pack(side="left", padx=(0, 28), pady=(4, 0))

        self.nav_buttons = {}
        for name, label in (("3 · CV & Export", "CV & Export"),
                            ("2 · Profile", "Profile"),
                            ("1 · Candidates", "Candidates")):
            item = ctk.CTkFrame(navbar, fg_color="transparent")
            item.pack(side="right", fill="y", padx=(0, 14))
            b = ctk.CTkButton(item, text=label, height=50, width=10,
                              corner_radius=0, fg_color="transparent",
                              hover_color=SURFACE, text_color=MUTED, font=FONT,
                              command=lambda n=name: self.show_tab(n))
            b.pack(side="top", fill="both", expand=True, padx=10)
            underline = ctk.CTkFrame(item, fg_color="transparent", height=3, corner_radius=0)
            underline.pack(side="bottom", fill="x", padx=10)
            self.nav_buttons[name] = (b, underline)

        # LibreOffice warning banner — shown only when soffice is missing
        if not _find_soffice():
            banner = ctk.CTkFrame(main, fg_color="#fef3e2", corner_radius=0, height=34)
            banner.pack(fill="x")
            banner.pack_propagate(False)
            ctk.CTkLabel(
                banner,
                text="⚠  LibreOffice is not installed — PDF export will have reduced quality.  "
                     "Click the wrench icon (⚙) in the sidebar to download it.",
                font=FONT_SM, text_color="#a06b1a",
            ).pack(expand=True)

        self.tabs = ctk.CTkTabview(main, fg_color="transparent")
        self.tabs.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        self.tab_upload = self.tabs.add("1 · Candidates")
        self.tab_profile = self.tabs.add("2 · Profile")
        self.tab_cv = self.tabs.add("3 · CV & Export")
        self.tabs._segmented_button.grid_forget()   # navbar drives the tabs now
        self.show_tab("1 · Candidates")

        self.status = ctk.CTkLabel(main, text="Ready", font=FONT_SM, text_color=MUTED)
        self.status.pack(pady=(0, 8))

        self._build_upload()
        self._build_profile()
        self._build_cv()

    def reconnect(self):
        """Called by the shell's global reload button — resets the Sheets singleton."""
        global _sheets_store_instance, _sheets_store_checked
        _sheets_store_instance = None
        _sheets_store_checked  = False
        store = _get_sheets_store()
        if store:
            self.set_status("Google Sheets reconnected ✓", GREEN)
        else:
            self.set_status("Google Sheets not available", RED)

    def set_status(self, text, color=MUTED):
        self.status.configure(text=text, text_color=color)

    def show_tab(self, name):
        self.tabs.set(name)
        for n, (b, underline) in self.nav_buttons.items():
            active = n == name
            b.configure(text_color=WHITE if active else MUTED,
                        font=FONT_BOLD if active else FONT)
            underline.configure(fg_color=GOLD if active else "transparent")

    # ── tab 1: candidates ──
    def _build_upload(self):
        f = centered_column(self.tab_upload, 870)
        top = ctk.CTkFrame(f, fg_color="transparent")
        top.pack(fill="x", padx=24, pady=(14, 4))
        ctk.CTkLabel(top, text="👤  Candidates", font=FONT_H1, text_color=WHITE).pack(side="left")
        ctk.CTkButton(top, text="＋  New Candidate", fg_color=GOLD, hover_color=GOLD_HV,
                      text_color="#ffffff", font=FONT_BOLD, height=38,
                      command=self.pick_file).pack(side="right")

        search_row = ctk.CTkFrame(f, fg_color="transparent")
        search_row.pack(fill="x", padx=24, pady=(6, 4))
        self.search_var = ctk.StringVar()
        self.search_var.trace_add("write", lambda *_: self.refresh_profile_list())
        ctk.CTkEntry(search_row, textvariable=self.search_var, height=36,
                     placeholder_text="Search by name and job title",
                     fg_color=CARD, border_color=HAIR, border_width=1,
                     corner_radius=10, text_color=WHITE).pack(fill="x")

        self.profile_list = ctk.CTkScrollableFrame(f, fg_color=CARD, border_color=HAIR,
                                                   border_width=1, corner_radius=14)
        self.profile_list.pack(fill="both", expand=True, padx=24, pady=(6, 16))
        self.refresh_profile_list()

    @staticmethod
    def _chip_color(text):
        i = sum(ord(c) for c in text) % len(CHIP_COLORS)
        return CHIP_COLORS[i], CHIP_TEXT[i]

    def refresh_profile_list(self):
        for w in self.profile_list.winfo_children():
            w.destroy()
        lbl = ctk.CTkLabel(self.profile_list, text="Loading candidates…",
                            font=FONT_SM, text_color=MUTED)
        lbl.pack(pady=16)
        self.profile_list.update_idletasks()

        def _fetch():
            try:
                summaries = list_profiles()
            except Exception as e:
                self.after(0, lambda: lbl.configure(text=f"Could not load profiles: {e}"))
                return
            self.after(0, lambda s=summaries: _render(s))

        def _render(summaries):
            lbl.destroy()
            query = (self.search_var.get() if hasattr(self, "search_var") else "").lower().strip()
            shown = 0
            hover_group = {"current": None}
            for s in summaries:
                profile_id = s.get("profile_id", "")
                name       = s.get("name") or "Unnamed"
                title      = s.get("job_title") or ""
                parsed     = s.get("parsed_date") or ""
                if query and query not in name.lower() and query not in title.lower():
                    continue
                shown += 1
                row = ctk.CTkFrame(self.profile_list, fg_color="transparent", corner_radius=8)
                row.pack(fill="x", pady=1)
                left = ctk.CTkFrame(row, fg_color="transparent")
                left.pack(side="left", fill="x", expand=True, padx=12, pady=8)
                ctk.CTkLabel(left, text=name, font=FONT_BOLD, text_color=WHITE,
                             anchor="w").pack(anchor="w")
                bits = []
                if s.get("work_count"):
                    bits.append(f"{s['work_count']} jobs")
                if s.get("edu_count"):
                    bits.append(f"{s['edu_count']} edu")
                if bits:
                    ctk.CTkLabel(left, text="  ·  ".join(bits), font=FONT_SM,
                                 text_color=MUTED, anchor="w").pack(anchor="w")
                if title:
                    bg, fg_ = self._chip_color(title)
                    ctk.CTkLabel(row, text="  " + title[:38] + "  ", font=FONT_SM,
                                 text_color=fg_, fg_color=bg, corner_radius=8).pack(side="left", padx=10)
                ctk.CTkButton(row, text="Load", width=58, fg_color=SURFACE, hover_color=HAIR,
                              border_color=HAIR, border_width=1, text_color=GOLD,
                              command=lambda pid=profile_id: self.load_saved(pid)).pack(side="right", padx=12)
                raw_link = s.get("raw_cv_link", "")
                if raw_link:
                    ctk.CTkButton(row, text="📄 CV", width=52, fg_color=SURFACE, hover_color=HAIR,
                                  border_color=HAIR, border_width=1, text_color=MUTED,
                                  font=FONT_SM,
                                  command=lambda url=raw_link: __import__("webbrowser").open(url)
                                  ).pack(side="right", padx=4)
                if parsed:
                    ctk.CTkLabel(row, text=parsed[:16], font=FONT_SM,
                                 text_color=MUTED).pack(side="right", padx=10)
                for wdg in (row, left):
                    wdg.bind("<Button-1>", lambda _e, pid=profile_id: self.load_saved(pid))
                bind_hover(row, row, SURFACE, "transparent", group=hover_group)
                ctk.CTkFrame(self.profile_list, fg_color=HAIR, height=1).pack(fill="x", padx=8)
            if not shown:
                ctk.CTkLabel(self.profile_list,
                             text="No candidates found" if query else
                                  "No candidates yet — click ＋ New Candidate to parse a CV",
                             font=FONT_SM, text_color=MUTED).pack(pady=16)

        threading.Thread(target=_fetch, daemon=True).start()

    def load_saved(self, profile_ref):
        try:
            profile, cv = load_profile(profile_ref)
        except Exception as e:
            messagebox.showerror("Load failed", str(e))
            return
        self.profile_path = profile_ref
        self.profile_form.populate(profile)
        if cv:
            self.cv_form.populate(cv)
        else:
            self.cv_form.populated = False
        self.show_tab("2 · Profile")
        label = os.path.basename(str(profile_ref)) if os.sep in str(profile_ref) \
                else str(profile_ref).split("_")[0]
        self.set_status(f"Loaded {label}", GREEN)

    def pick_file(self):
        path = filedialog.askopenfilename(
            title="Choose a CV",
            filetypes=[("CV files", "*.docx *.doc *.pdf"), ("All files", "*.*")])
        if not path:
            return
        self._show_parse_overlay(os.path.basename(path))
        self.set_status(f"Parsing {os.path.basename(path)}…", ORANGE)
        threading.Thread(target=self._parse_worker, args=(path,), daemon=True).start()

    # ── parse loading overlay ──
    def _show_parse_overlay(self, filename):
        self._overlay = ctk.CTkFrame(self, fg_color=BG)
        self._overlay.place(relx=0, rely=0, relwidth=1, relheight=1)
        box = ctk.CTkFrame(self._overlay, fg_color=CARD, corner_radius=16,
                           border_width=1, border_color=HAIR)
        box.place(relx=0.5, rely=0.45, anchor="center")
        inner = ctk.CTkFrame(box, fg_color="transparent")
        inner.pack(padx=60, pady=40)
        self._spinner_lbl = ctk.CTkLabel(inner, text="◐", font=("Arial", 44), text_color=GOLD)
        self._spinner_lbl.pack()
        ctk.CTkLabel(inner, text="Reading the CV", font=FONT_H1, text_color=WHITE).pack(pady=(12, 2))
        ctk.CTkLabel(inner, text=filename, font=FONT_SM, text_color=MUTED).pack()
        self._overlay_step = ctk.CTkLabel(inner, text="Extracting text…", font=FONT_SM,
                                          text_color=GOLD)
        self._overlay_step.pack(pady=(14, 0))
        self._spinner_frames = ["◐", "◓", "◑", "◒"]
        self._spinner_i = 0
        self._overlay_steps = ["Extracting text…", "Identifying candidate details…",
                               "Structuring work history…", "Collecting skills & education…",
                               "Nearly there…"]
        self._overlay_step_i = 0
        self._overlay_tick = 0
        self._animate_overlay()

    def _animate_overlay(self):
        if not getattr(self, "_overlay", None) or not self._overlay.winfo_exists():
            return
        self._spinner_i = (self._spinner_i + 1) % len(self._spinner_frames)
        self._spinner_lbl.configure(text=self._spinner_frames[self._spinner_i])
        self._overlay_tick += 1
        if self._overlay_tick % 14 == 0 and self._overlay_step_i < len(self._overlay_steps) - 1:
            self._overlay_step_i += 1
            self._overlay_step.configure(text=self._overlay_steps[self._overlay_step_i])
        self.after(180, self._animate_overlay)

    def _hide_parse_overlay(self):
        if getattr(self, "_overlay", None) and self._overlay.winfo_exists():
            self._overlay.destroy()
        self._overlay = None

    def _parse_worker(self, path):
        try:
            profile = parse_cv(path)
        except Exception as e:
            traceback.print_exc()
            err = str(e)
            self.after(0, lambda: (self._hide_parse_overlay(),
                                   self.set_status("Parse failed", RED),
                                   messagebox.showerror("Parse failed", err)))
            return

        self.after(0, lambda: self._parse_done(profile))

    def _parse_done(self, profile):
        self._hide_parse_overlay()
        self.profile_path = None
        self._pending_raw_cv_link = ""
        self.profile_form.populate(profile)
        self.cv_form.populated = False
        self.show_tab("2 · Profile")
        self.set_status(f"Parsed — {len(profile.work_history)} jobs, "
                        f"{len(profile.education)} education entries. Review the profile.", GREEN)

    # ── tab 2: profile (source of truth) ──
    def _build_profile(self):
        f = centered_column(self.tab_profile, 870)
        self.profile_form = ProfileForm(self, f)
        bar = ctk.CTkFrame(f, fg_color="transparent")
        bar.pack(fill="x", pady=8)
        ctk.CTkLabel(bar, text="The profile is the candidate's master record — kept separate from the CV.",
                     font=FONT_SM, text_color=MUTED).pack(side="left", padx=12)
        ctk.CTkButton(bar, text="Continue to CV →", fg_color=GOLD, hover_color=GOLD_HV,
                      text_color="#ffffff", font=FONT_BOLD,
                      command=self.go_cv).pack(side="right", padx=10)
        ctk.CTkButton(bar, text="Save Profile", fg_color=CARD, hover_color=SURFACE,
                      border_color=HAIR, border_width=1, text_color=GOLD,
                      command=self.do_save_profile).pack(side="right", padx=6)

    def do_save_profile(self):
        profile = self.profile_form.collect()
        if not profile:
            messagebox.showinfo("Nothing to save", "Parse or load a CV first.")
            return
        cv = self.cv_form.collect()
        self.set_status("Saving…", ORANGE)
        raw_link = self._pending_raw_cv_link
        def _save():
            try:
                ref = save_profile(profile, cv, existing_path=self.profile_path,
                                   raw_cv_link=raw_link)
                self.after(0, lambda: self._on_save_done(ref))
            except Exception as e:
                self.after(0, lambda: (self.set_status("Save failed", RED),
                                       messagebox.showerror("Save failed", str(e))))
        threading.Thread(target=_save, daemon=True).start()

    def _on_save_done(self, ref):
        self.profile_path = ref
        label = os.path.basename(str(ref)) if os.sep in str(ref) else str(ref).split("_")[0]
        self.set_status(f"Saved — {label}", GREEN)
        self.refresh_profile_list()

    def go_cv(self):
        profile = self.profile_form.collect()
        if not profile:
            messagebox.showinfo("No profile", "Parse or load a CV first.")
            return
        if self.cv_form.populated:
            if not messagebox.askyesno(
                    "Refresh CV fields?",
                    "Copy the profile into the CV fields again?\n"
                    "This overwrites any CV-only edits you've made."):
                self.show_tab("3 · CV & Export")
                return
        self.cv_form.populate(profile)
        self.show_tab("3 · CV & Export")
        self.set_status("CV fields initialised from profile — edits here affect the CV only", BLUE)
        self.schedule_preview(800)

    # ── tab 3: CV & export ──
    def _build_cv(self):
        f = self.tab_cv
        left = ctk.CTkFrame(f, fg_color="transparent")
        left.pack(side="left", fill="both", expand=True, padx=(4, 4), pady=4)
        right = ctk.CTkFrame(f, fg_color=CARD, border_color=HAIR, border_width=1, corner_radius=14)
        right.pack(side="right", fill="both", expand=True, padx=(4, 6), pady=8)

        self.cv_form = ProfileForm(self, left, mode="cv", on_change=self.schedule_preview)
        bar = ctk.CTkFrame(left, fg_color="transparent")
        bar.pack(fill="x", pady=6)
        ctk.CTkLabel(bar, text="Edits here change the CV only — the profile is untouched.",
                     font=FONT_SM, text_color=MUTED).pack(side="left", padx=12)

        # ── right pane: controls (bottom) + preview (fills the rest) ──
        controls = ctk.CTkFrame(right, fg_color="transparent")
        controls.pack(side="bottom", fill="x", pady=(2, 10))
        rowA = ctk.CTkFrame(controls, fg_color="transparent")
        rowA.pack(pady=2)
        ctk.CTkButton(rowA, text="◀", width=32, fg_color=SURFACE, hover_color=HAIR,
                      border_color=HAIR, border_width=1, text_color=WHITE,
                      command=lambda: self.preview_nav(-1)).pack(side="left", padx=3)
        self.preview_page_lbl = ctk.CTkLabel(rowA, text="– / –", font=FONT_SM, text_color=MUTED)
        self.preview_page_lbl.pack(side="left", padx=6)
        ctk.CTkButton(rowA, text="▶", width=32, fg_color=SURFACE, hover_color=HAIR,
                      border_color=HAIR, border_width=1, text_color=WHITE,
                      command=lambda: self.preview_nav(1)).pack(side="left", padx=3)
        self.preview_status = ctk.CTkLabel(rowA, text="", font=FONT_SM, text_color=MUTED)
        self.preview_status.pack(side="left", padx=10)

        rowB = ctk.CTkFrame(controls, fg_color="transparent")
        rowB.pack(pady=4)
        ensure_template()
        tpls = list_templates() or [DEFAULT_TEMPLATE]
        current = CONFIG["formatting"].get("template", DEFAULT_TEMPLATE)
        self.template_var = ctk.StringVar(value=current if current in tpls else tpls[0])
        self.template_menu = ctk.CTkOptionMenu(rowB, variable=self.template_var, values=tpls,
                                               fg_color=SURFACE, button_color=SURFACE,
                                               button_hover_color=HAIR, text_color=WHITE, width=190,
                                               command=lambda _v: self.schedule_preview(300))
        self.template_menu.pack(side="left", padx=4)
        ctk.CTkButton(rowB, text="Edit Template", width=104, fg_color=SURFACE, hover_color=HAIR,
                      border_color=HAIR, border_width=1, text_color=WHITE,
                      command=self.edit_template).pack(side="left", padx=3)
        ctk.CTkButton(rowB, text="Export DOCX", width=110, fg_color=SURFACE, hover_color=HAIR,
                      border_color=HAIR, border_width=1, text_color=WHITE,
                      command=self.do_generate_docx).pack(side="left", padx=3)
        ctk.CTkButton(rowB, text="Export PDF", width=120, fg_color=GOLD, hover_color=GOLD_HV,
                      text_color="#ffffff", font=FONT_BOLD,
                      command=self.do_generate).pack(side="left", padx=6)

        backdrop = ctk.CTkFrame(right, fg_color=SURFACE, corner_radius=10)
        backdrop.pack(side="top", fill="both", expand=True, padx=10, pady=(10, 2))
        self.preview_label = ctk.CTkLabel(backdrop, text="The CV preview renders here automatically\nas you edit",
                                          font=FONT_SM, text_color=MUTED)
        self.preview_label.pack(fill="both", expand=True, padx=8, pady=8)
        backdrop.bind("<Configure>", self._preview_rescale)
        self._preview_pane = backdrop

    def edit_template(self):
        path = template_path(self.template_var.get())
        if not os.path.exists(path):
            ensure_template()
        if sys.platform == "darwin":
            os.system(f'open "{path}"')
        elif sys.platform == "win32":
            os.startfile(path)
        self.set_status("Template opened — save in Word, the preview refreshes on next edit", BLUE)

    # ── auto preview (debounced) ──
    def schedule_preview(self, delay_ms=1600):
        if not self.cv_form.populated:
            return
        if self._preview_after:
            self.after_cancel(self._preview_after)
        self._preview_after = self.after(delay_ms, self.update_preview)

    def update_preview(self):
        self._preview_after = None
        if not self.cv_form.populated:
            return
        if self._preview_busy:
            self._preview_dirty = True
            return
        cv = self.cv_form.collect()
        self._preview_busy = True
        self.preview_status.configure(text="rendering…", text_color=ORANGE)
        threading.Thread(target=self._preview_worker, args=(cv,), daemon=True).start()

    def _preview_worker(self, cv):
        import tempfile
        try:
            import fitz  # pymupdf — bundled, no external binary needed
        except ImportError:
            self.after(0, lambda: self.preview_status.configure(
                text="pip install pymupdf required for preview", text_color=RED))
            self._preview_busy = False
            return
        try:
            tmp = tempfile.mkdtemp()
            _, pdf = generate_cv(cv, CONFIG["formatting"], photo_path=self.photo_path,
                                 template_name=self.template_var.get(), out_dir=tmp)
            doc = fitz.open(pdf)
            from PIL import Image
            pages = []
            for page in doc:
                mat = fitz.Matrix(150 / 72, 150 / 72)  # 150 dpi
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                pages.append(img)
            doc.close()
            if not pages:
                raise RuntimeError("No pages rendered")
            self.preview_pil = pages
            self.preview_idx = min(self.preview_idx, len(self.preview_pil) - 1)
            self.after(0, self._preview_show)
        except Exception as e:
            traceback.print_exc()
            err = str(e)
            self.after(0, lambda: self.preview_status.configure(
                text=f"failed: {err[:60]}", text_color=RED))
        finally:
            self._preview_busy = False
            if self._preview_dirty:
                self._preview_dirty = False
                self.after(0, lambda: self.schedule_preview(400))

    def _preview_show(self):
        if not self.preview_pil:
            return
        from PIL import Image, ImageOps
        pane_w = max(self._preview_pane.winfo_width() - 28, 300)
        pane_h = max(self._preview_pane.winfo_height() - 28, 300)
        im = self.preview_pil[self.preview_idx]
        scale = min(pane_w / im.width, pane_h / im.height)
        size = (max(int(im.width * scale), 1), max(int(im.height * scale), 1))
        # high-quality downscale + hairline page edge so the paper stands out
        sharp = im.resize(size, Image.LANCZOS)
        sharp = ImageOps.expand(sharp, border=1, fill="#c6c0b2")
        img = ctk.CTkImage(light_image=sharp, size=(size[0] + 2, size[1] + 2))
        self._preview_img_ref = img
        self.preview_label.configure(image=img, text="")
        self.preview_page_lbl.configure(text=f"{self.preview_idx + 1} / {len(self.preview_pil)}")
        self.preview_status.configure(text="up to date", text_color=GREEN)

    def _preview_rescale(self, _e=None):
        if self.preview_pil:
            if getattr(self, "_rescale_after", None):
                self.after_cancel(self._rescale_after)
            self._rescale_after = self.after(150, self._preview_show)

    def preview_nav(self, delta):
        if not self.preview_pil:
            return
        self.preview_idx = max(0, min(len(self.preview_pil) - 1, self.preview_idx + delta))
        self._preview_show()

    # ── export ──
    def do_generate_docx(self):
        """Export the formatted CV as a DOCX file (always full quality — no Word required)."""
        cv = self.cv_form.collect()
        if not cv:
            messagebox.showinfo("No CV", "Go through Profile → Continue to CV first.")
            return
        job      = (cv.job_title or "CV").strip()
        county   = (cv.county or "").strip()
        initials = "".join(w[0].upper() for w in (cv.name or "").split() if w)
        parts    = [job]
        if county:  parts.append(county)
        if initials: parts.append(initials)
        default_name = " - ".join(parts) + ".docx"
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        save_path = filedialog.asksaveasfilename(
            title="Save formatted CV as DOCX…",
            initialdir=OUTPUT_DIR, initialfile=default_name,
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
            parent=self,
        )
        if not save_path:
            return
        self.set_status("Exporting DOCX…", ORANGE)
        threading.Thread(target=self._generate_docx_worker, args=(cv, save_path), daemon=True).start()

    def _generate_docx_worker(self, cv, save_path):
        try:
            import shutil, tempfile
            tmp_dir  = tempfile.mkdtemp()
            docx_tmp, _ = generate_cv(cv, CONFIG["formatting"], photo_path=self.photo_path,
                                      template_name=self.template_var.get(), out_dir=tmp_dir,
                                      keep_docx=True)
            if docx_tmp and os.path.exists(docx_tmp):
                shutil.move(docx_tmp, save_path)
            else:
                raise RuntimeError("DOCX was not created")
        except Exception as e:
            traceback.print_exc()
            err = str(e)
            self.after(0, lambda: (self.set_status("Export failed", RED),
                                   messagebox.showerror("Export failed", err)))
            return
        def done():
            self.set_status(f"Saved → {os.path.basename(save_path)}", GREEN)
            if sys.platform == "darwin":
                os.system(f'open -R "{save_path}"')
            elif sys.platform == "win32":
                os.startfile(os.path.dirname(save_path))
        self.after(0, done)

    def do_generate(self):
        cv = self.cv_form.collect()
        if not cv:
            messagebox.showinfo("No CV", "Go through Profile → Continue to CV first.")
            return

        # Build default filename: "Job Title - County - Initials.pdf"
        job   = (cv.job_title or "CV").strip()
        county = (cv.county or "").strip()
        initials = "".join(w[0].upper() for w in (cv.name or "").split() if w)
        parts = [job]
        if county:
            parts.append(county)
        if initials:
            parts.append(initials)
        default_name = " - ".join(parts) + ".pdf"

        os.makedirs(OUTPUT_DIR, exist_ok=True)
        save_path = filedialog.asksaveasfilename(
            title="Save formatted CV as…",
            initialdir=OUTPUT_DIR,
            initialfile=default_name,
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            parent=self,
        )
        if not save_path:
            return

        self.set_status("Exporting PDF…", ORANGE)
        threading.Thread(target=self._generate_worker, args=(cv, save_path), daemon=True).start()

    def _generate_worker(self, cv, save_path):
        try:
            import shutil, tempfile
            tmp_dir = tempfile.mkdtemp()
            _, pdf_tmp = generate_cv(cv, CONFIG["formatting"], photo_path=self.photo_path,
                                     template_name=self.template_var.get(), out_dir=tmp_dir)
            shutil.move(pdf_tmp, save_path)
        except Exception as e:
            traceback.print_exc()
            err = str(e)
            self.after(0, lambda: (self.set_status("Export failed", RED),
                                   messagebox.showerror("Export failed", err)))
            return

        def done():
            folder = os.path.dirname(save_path)
            self.set_status(f"Saved → {os.path.basename(save_path)}", GREEN)
            if sys.platform == "darwin":
                os.system(f'open "{folder}"')
            elif sys.platform == "win32":
                os.startfile(folder)
        self.after(0, done)

def _standalone():
    ctk.set_appearance_mode("light")
    ensure_template()
    root = ctk.CTk()
    root.title("CV Parse & Format Tool  V0.1")
    root.geometry("1320x840")
    root.configure(fg_color=BG)
    tool = CVParseFormatTool(root)
    tool.pack(fill="both", expand=True)
    root.mainloop()


if __name__ == "__main__":
    _standalone()
